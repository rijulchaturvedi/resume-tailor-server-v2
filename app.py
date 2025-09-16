from flask import Flask, request, send_file, jsonify, make_response
from flask_cors import CORS
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.oxml.ns import qn
import io, os, json, re, logging
from typing import List, Dict, Tuple, Optional
from concurrent.futures import ThreadPoolExecutor, TimeoutError as FuturesTimeout
from datetime import datetime

# ----------------------------
# App & CORS
# ----------------------------
app = Flask(__name__)
CORS(app, resources={r"/tailor": {"origins": "chrome-extension://*"}})
app.logger.setLevel(logging.INFO)

# ----------------------------
# OpenAI config
# ----------------------------
MODEL = os.getenv("OPENAI_MODEL", "gpt-4o-mini")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
EXECUTOR = ThreadPoolExecutor(max_workers=int(os.getenv("OAI_THREADS", "2")))
OAI_ENABLED = os.getenv("USE_OPENAI", "1").strip().lower() not in ("0", "false", "no")
OAI_BUDGET = float(os.getenv("OAI_BUDGET", "60"))  # Increased timeout
OAI_CLIENT_TIMEOUT = float(os.getenv("OAI_CLIENT_TIMEOUT", "45"))  # Increased timeout
USE_RESPONSES_API = os.getenv("USE_RESPONSES_API", "0").strip().lower() in ("1","true","yes")

def _env_bool(name: str, default: bool=False) -> bool:
    v = os.getenv(name)
    if v is None: return default
    return v.strip().lower() in ("1","true","yes","on","y")

# Behavior toggles
SHOW_KPI_PLACEHOLDER = _env_bool("SHOW_KPI_PLACEHOLDER", False)
BULLETS_STRICT_REPLACE = _env_bool("BULLETS_STRICT_REPLACE", True)

try:
    from openai import OpenAI
    _openai_available = True
except Exception:
    _openai_available = False

# ----------------------------
# Text helpers
# ----------------------------
KNOWN_HEADINGS = {
    "PROFESSIONAL SUMMARY","SUMMARY","EXECUTIVE SUMMARY",
    "EXPERIENCE","WORK EXPERIENCE","PROFESSIONAL EXPERIENCE",
    "SKILLS","CORE SKILLS","TECHNICAL SKILLS","SKILLS & TOOLS","SKILLS AND TOOLS","CORE COMPETENCIES",
    "EDUCATION","PROJECTS","TECHNICAL PROJECTS","CERTIFICATIONS","PUBLICATIONS","ACHIEVEMENTS","ACADEMIC PROJECTS"
}

def humanize_text(text: str) -> str:
    """Remove AI-style dashes and make text more natural"""
    if not text: return ""
    # Replace em/en dashes with simple hyphens or commas
    text = text.replace("–", ", ").replace("—", ", ")
    # Replace smart quotes with regular quotes
    text = re.sub(r"[\u201c\u201d]", '"', text)
    text = re.sub(r"[\u2018\u2019]", "'", text)
    # Remove excessive punctuation patterns common in AI text
    text = re.sub(r'\s*,\s*,+', ',', text)
    text = re.sub(r'\s*;\s*;+', ';', text)
    return text.strip()

def sanitize(text: str) -> str:
    if not text: return ""
    text = text.replace("\u00A0", " ").replace("\t", " ")
    text = humanize_text(text)
    text = re.sub(r"\s{2,}", " ", text)
    return text.strip()

def _canon(s: str) -> str:
    s = sanitize(s).lower()
    s = re.sub(r"[^\w\s]", " ", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s

def _norm_heading(text: str) -> str:
    t = sanitize(text).upper()
    t = t.replace("&","AND")
    t = re.sub(r"[^A-Z0-9 ]+", " ", t)
    t = re.sub(r"\s{2,}", " ", t).strip()
    return t

def is_major_heading(text: str) -> bool:
    t = _norm_heading(text)
    return (not t) or t in KNOWN_HEADINGS or (t.isupper() and len(t.split()) <= 4)

# ----------------------------
# DOCX utils
# ----------------------------
def ensure_docx(doc_or_bytes):
    try:
        if hasattr(doc_or_bytes, "paragraphs"):
            return doc_or_bytes
        if hasattr(doc_or_bytes, "read"):
            data = doc_or_bytes.read()
        else:
            data = doc_or_bytes
        return Document(io.BytesIO(data))
    except Exception as e:
        raise RuntimeError(f"Failed to open DOCX: {e}")

def delete_range(doc: Document, start: int, end: int):
    """Delete paragraphs from start to end (inclusive)"""
    for i in range(end, start-1, -1):
        if i < len(doc.paragraphs):
            p = doc.paragraphs[i]._element
            p.getparent().remove(p)

def set_paragraph_font(paragraph: Paragraph, font_name: str = "Times New Roman", font_size: int = 9):
    """Set font for all runs in a paragraph"""
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        # Set the font for both ASCII and East Asian text
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)

def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: Optional[str] = None) -> Paragraph:
    """Insert a new paragraph after the given paragraph with proper formatting"""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    run = None
    if text:
        run = new_para.add_run(text)
        # Set font to Times New Roman, size 9
        run.font.name = "Times New Roman"
        run.font.size = Pt(9)
        # Ensure font is set for both ASCII and East Asian text
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:ascii'), "Times New Roman")
        rFonts.set(qn('w:hAnsi'), "Times New Roman")
    if style:
        try:
            new_para.style = style
        except KeyError:
            if run and style.lower().startswith("list") and not run.text.strip().startswith("•"):
                run.text = f"• {run.text}"
    return new_para

def find_section_bounds(doc: Document, titles: List[str]) -> Optional[Tuple[int,int]]:
    titles_up = {_norm_heading(t) for t in titles}
    start = None
    for i,p in enumerate(doc.paragraphs):
        if _norm_heading(p.text) in titles_up:
            start = i
            break
    if start is None:
        return None
    end = len(doc.paragraphs)-1
    for j in range(start+1, len(doc.paragraphs)):
        if is_major_heading(doc.paragraphs[j].text):
            end = j-1
            break
    return (start, end)

def find_all_section_bounds(doc: Document, titles: List[str]) -> List[Tuple[int,int]]:
    spans = []
    i = 0
    titles_up = {_norm_heading(t) for t in titles}
    while i < len(doc.paragraphs):
        if _norm_heading(doc.paragraphs[i].text) in titles_up:
            s = i
            e = len(doc.paragraphs)-1
            for j in range(i+1, len(doc.paragraphs)):
                if is_major_heading(doc.paragraphs[j].text):
                    e = j-1
                    break
            spans.append((s,e))
            i = e + 1
        else:
            i += 1
    return spans

def find_anchors_by_exact_headers(doc: Document, exact_headers: Dict[str,str]) -> List[Tuple[str,int]]:
    pairs = [(sanitize(k), v) for k, v in exact_headers.items()]
    found = []
    used = set()
    for comp, header in pairs:
        idx = None
        for i, p in enumerate(doc.paragraphs):
            if sanitize(p.text) == sanitize(header):
                idx = i; break
        if idx is None:
            for i, p in enumerate(doc.paragraphs):
                if _canon(header) in _canon(p.text):
                    idx = i; break
        if idx is not None and idx not in used:
            found.append((comp, idx)); used.add(idx)
        else:
            app.logger.warning("Header NOT found for %s -> '%s'", comp, header)
    found.sort(key=lambda x: x[1])
    return found

def is_bullet_point(text: str) -> bool:
    """Check if a paragraph is a bullet point"""
    text = text.strip()
    # Check for common bullet markers including the special bullet character
    bullet_markers = ['•', '◦', '○', '▪', '▫', '▪', '▫', '-', '*', '>', '→', '·']
    for marker in bullet_markers:
        if text.startswith(marker):
            return True
    # Also check if it starts with a number followed by period or parenthesis (numbered list)
    if re.match(r'^\d+[.)]\s', text):
        return True
    return False

def find_role_section_bounds(doc: Document, header_idx: int) -> Tuple[int, int]:
    """Find bounds of ALL content under a role header, including all bullets"""
    content_start = header_idx + 1
    content_end = len(doc.paragraphs) - 1
    
    # Continue until we hit a major heading or another role header
    for i in range(header_idx + 1, len(doc.paragraphs)):
        para_text = sanitize(doc.paragraphs[i].text)
        
        # Skip empty paragraphs - they don't mark the end
        if not para_text:
            continue
        
        # Check if it's a major heading (EDUCATION, CERTIFICATIONS, etc.)
        if is_major_heading(para_text):
            content_end = i - 1
            break
        
        # Check if it's another role header (has bold text and role keywords)
        # But be careful not to stop at bullets
        if i > header_idx + 1:  # Don't check the very first line after header
            para = doc.paragraphs[i]
            # Check if paragraph has bold formatting (typical of role headers)
            has_bold = any(run.bold for run in para.runs)
            
            # Only consider it a new role if it has bold AND looks like a role title
            if has_bold and _is_role_header(para_text, doc, i):
                # Double-check it's not a bullet by checking for bullet markers
                if not any(para_text.strip().startswith(marker) for marker in ['•', '·', '-', '*', '▪', '▫']):
                    content_end = i - 1
                    break
    
    # Extend the end to include any bullets that might follow
    # Look for continuous content that looks like bullets
    actual_end = content_end
    for i in range(content_end + 1, min(content_end + 10, len(doc.paragraphs))):
        para_text = doc.paragraphs[i].text.strip()
        # If it's a bullet-like line, include it
        if para_text and (para_text.startswith('-') or para_text.startswith('•') or 
                         para_text.startswith('·') or len(para_text) > 40):
            actual_end = i
        # Stop if we hit a clear section header
        elif is_major_heading(para_text):
            break
        # Stop if we hit another role header with bold
        elif any(run.bold for run in doc.paragraphs[i].runs):
            break
    
    return content_start, actual_end

def _is_role_header(text: str, doc: Document = None, idx: int = None) -> bool:
    """Enhanced check if paragraph is a role header"""
    text = sanitize(text).lower()
    if not text: return False
    
    # Check for bold formatting (strong indicator of headers)
    if doc and idx is not None and idx < len(doc.paragraphs):
        para = doc.paragraphs[idx]
        if para.runs and any(run.bold for run in para.runs):
            # If it has bold text and contains role/company indicators
            role_indicators = [
                "analyst", "manager", "engineer", "developer", "consultant", 
                "director", "specialist", "coordinator", "lead", "senior", "intern",
                "advisor", "architect", "administrator", "officer"
            ]
            if any(keyword in text for keyword in role_indicators):
                return True
    
    # Patterns that indicate role headers
    role_indicators = [
        "analyst", "manager", "engineer", "developer", "consultant", 
        "director", "specialist", "coordinator", "lead", "senior", "intern"
    ]
    
    # Date patterns
    has_date = bool(re.search(r'\b(19|20)\d{2}\b|present|current', text))
    
    # Company/location patterns
    has_location = bool(re.search(r',\s*[A-Z]{2}\b|,\s*(ny|ca|tx|wa|il|ma)\b', text, re.IGNORECASE))
    
    has_role = any(keyword in text for keyword in role_indicators)
    
    # Strong indicators: role + date or role + location
    if has_role and (has_date or has_location):
        return True
    
    # Check for pattern like "Title, Company, Location Date"
    if "," in text and (has_date or has_location):
        return True
    
    return False

# ----------------------------
# Professional Summary (Replace)
# ----------------------------
def replace_summary(doc: Document, summary: str):
    """Replace existing summary completely"""
    sec = find_section_bounds(doc, ["PROFESSIONAL SUMMARY","SUMMARY","EXECUTIVE SUMMARY"])
    if sec:
        s, e = sec
        # Delete all content under summary heading
        if e >= s+1:
            delete_range(doc, s+1, e)
        # Insert new summary with proper formatting
        new_para = insert_paragraph_after(doc.paragraphs[s], sanitize(summary))
        set_paragraph_font(new_para)
    else:
        # Add new summary section if not exists
        h = doc.add_paragraph()
        run = h.add_run("EXECUTIVE SUMMARY")
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(9)
        p = doc.add_paragraph(sanitize(summary))
        set_paragraph_font(p)

# ----------------------------
# Enhanced Skills Processing with Better Deduplication
# ----------------------------

def normalize_skill_for_comparison(skill: str) -> str:
    """Normalize skill names for better duplicate detection"""
    skill = sanitize(skill).lower()
    # Remove common prefixes/suffixes and variations
    skill = re.sub(r'\s*\([^)]*\)', '', skill)  # Remove parentheses content like "(BRD/FRD)"
    skill = re.sub(r'\s+', ' ', skill).strip()
    # Handle common variations
    skill = skill.replace('&', 'and')
    skill = skill.replace(' / ', '/')
    skill = skill.replace(' - ', '-')
    skill = skill.replace('aws ', '')
    skill = skill.replace('amazon ', '')
    return skill

def are_skills_similar(skill1: str, skill2: str) -> bool:
    """Check if two skills are similar enough to be considered duplicates"""
    norm1 = normalize_skill_for_comparison(skill1)
    norm2 = normalize_skill_for_comparison(skill2)
    
    # Exact match after normalization
    if norm1 == norm2:
        return True
    
    # Check if one is contained in the other (for cases like "SQL" and "SQL Queries")
    if (norm1 in norm2 and len(norm1) > 3) or (norm2 in norm1 and len(norm2) > 3):
        return True
    
    # Check for AWS services specifically
    if 'lambda' in norm1 and 'lambda' in norm2:
        return True
    if 's3' in norm1 and 's3' in norm2:
        return True
    if 'redshift' in norm1 and 'redshift' in norm2:
        return True
    
    return False

def parse_skills_section(doc: Document, s: int, e: int):
    """Parse existing skills maintaining structure and removing duplicates with better logic"""
    order = []
    mapping: Dict[str, List[str]] = {}
    current = None
    i = s + 1
    
    while i <= e and i < len(doc.paragraphs):
        line = sanitize(doc.paragraphs[i].text)
        if not line:
            i += 1
            continue
        
        # Check if this is a category header (bold text followed by colon)
        para = doc.paragraphs[i]
        has_bold = any(run.bold for run in para.runs) if para.runs else False
        
        # Category with colon (either standalone or inline)
        if ":" in line:
            parts = line.split(":", 1)
            if len(parts) == 2:
                head, items_part = parts
                category_key = _norm_heading(head.strip())
                
                if category_key not in mapping:
                    mapping[category_key] = []
                    order.append(category_key)
                
                current = category_key
                
                # Process items if they exist on the same line
                if items_part.strip():
                    items = [x.strip() for x in re.split(r'[,;]', items_part) if x.strip()]
                    mapping[category_key].extend(items)
        
        # Continuation line for current category (no colon, not bold)
        elif current and not has_bold:
            items = [x.strip() for x in re.split(r'[,;]', line) if x.strip()]
            mapping[current].extend(items)
        
        i += 1
    
    # Enhanced deduplication within each category
    for category_key, items in mapping.items():
        deduplicated = []
        
        for item in items:
            item_clean = item.strip().rstrip(',').rstrip('.')
            if len(item_clean) <= 2:  # Skip very short items
                continue
                
            # Check if this item is similar to any already added
            is_duplicate = False
            for existing in deduplicated:
                if are_skills_similar(item_clean, existing):
                    # Keep the longer/more descriptive version
                    if len(item_clean) > len(existing):
                        deduplicated.remove(existing)
                        deduplicated.append(item_clean)
                    is_duplicate = True
                    break
            
            if not is_duplicate:
                deduplicated.append(item_clean)
        
        mapping[category_key] = deduplicated
    
    return order, mapping

def rewrite_skills_section(doc: Document, s: int, e: int, order: List[str], mapping: Dict[str, List[str]]):
    """Rewrite skills section with proper bold headers and structured formatting"""
    if e >= s+1:
        delete_range(doc, s+1, e)
    
    anchor = doc.paragraphs[s]
    last = anchor
    
    for cat in order:
        # Format category name properly
        formatted_cat = cat.replace('_', ' ').title()
        formatted_cat = formatted_cat.replace(' And ', ' & ')
        formatted_cat = formatted_cat.replace('Ai ', 'AI ')  # Fix AI capitalization
        
        items = mapping.get(cat, [])
        
        if items:
            # Create category header paragraph with bold formatting
            header_text = f"{formatted_cat}:"
            header_para = insert_paragraph_after(last, header_text)
            set_paragraph_font(header_para)
            
            # Make the category header bold
            for run in header_para.runs:
                run.bold = True
            
            # Create items paragraph
            items_text = f"{', '.join(items)}"
            items_para = insert_paragraph_after(header_para, items_text)
            set_paragraph_font(items_para)
            
            last = items_para
    
    # Add some spacing after skills section
    spacing_para = insert_paragraph_after(last, "")
    set_paragraph_font(spacing_para)

def merge_skills(doc: Document, new_skills: Dict[str, list]):
    """Preserve existing skills and append new ones with enhanced deduplication"""
    if not new_skills:
        return
    
    spans = find_all_section_bounds(doc, ["SKILLS","CORE SKILLS","TECHNICAL SKILLS","SKILLS & TOOLS","SKILLS AND TOOLS","CORE COMPETENCIES"])
    
    if not spans:
        # Create skills section if not exists
        h = doc.add_paragraph()
        run = h.add_run("CORE COMPETENCIES")
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(9)
        spans = find_all_section_bounds(doc, ["CORE COMPETENCIES"])
    
    if not spans:
        return
    
    first_s = spans[0][0]
    last_e = spans[-1][1]
    
    # Parse existing skills
    order, mapping = parse_skills_section(doc, first_s, last_e)
    
    app.logger.info(f"Parsed existing skills: {list(mapping.keys())}")
    for cat, items in mapping.items():
        app.logger.info(f"  {cat}: {len(items)} items - {items[:3]}...")
    
    # Merge new skills with enhanced duplicate checking
    for cat, additions in new_skills.items():
        key = _norm_heading(cat)
        if key not in mapping:
            mapping[key] = []
            order.append(key)
        
        existing_skills = mapping[key]
        app.logger.info(f"Merging {len(additions or [])} new skills into {cat} (existing: {len(existing_skills)})")
        
        for new_skill in (additions or []):
            new_skill_clean = sanitize(new_skill)
            if not new_skill_clean or len(new_skill_clean) <= 2:
                continue
            
            # Check if this new skill is similar to any existing skill
            is_duplicate = False
            for existing_skill in existing_skills:
                if are_skills_similar(new_skill_clean, existing_skill):
                    # Keep the longer/more descriptive version
                    if len(new_skill_clean) > len(existing_skill):
                        existing_skills.remove(existing_skill)
                        existing_skills.append(new_skill_clean)
                        app.logger.info(f"  Replaced '{existing_skill}' with '{new_skill_clean}'")
                    else:
                        app.logger.info(f"  Skipped duplicate '{new_skill_clean}' (keeping '{existing_skill}')")
                    is_duplicate = True
                    break
            
            if not is_duplicate:
                existing_skills.append(new_skill_clean)
                app.logger.info(f"  Added new skill: '{new_skill_clean}'")
        
        mapping[key] = existing_skills
    
    # Rewrite section with merged content and proper formatting
    app.logger.info(f"Rewriting skills section with {len(order)} categories")
    rewrite_skills_section(doc, first_s, last_e, order, mapping)

# ----------------------------
# OpenAI Integration
# ----------------------------
def _get_client():
    if not OPENAI_API_KEY or not _openai_available or not OAI_ENABLED:
        return None
    try:
        return OpenAI(api_key=OPENAI_API_KEY, timeout=OAI_CLIENT_TIMEOUT, max_retries=0)
    except Exception as e:
        app.logger.exception("OpenAI client init failed: %s", e)
        return None

def _gpt_call(client, system, prompt) -> str:
    resp = client.chat.completions.create(
        model=MODEL,
        messages=[{"role":"system","content":system},{"role":"user","content":prompt}],
        temperature=0.3  # Lower temperature for consistency
    )
    return resp.choices[0].message.content.strip()

def gpt(prompt: str, system: str = "You are a professional resume writer. Write in a natural, human style without AI markers.") -> str:
    app.logger.info("=== GPT CALL START ===")
    app.logger.info(f"System: {system[:100]}...")
    app.logger.info(f"Prompt length: {len(prompt)} chars")
    app.logger.info(f"Prompt preview: {prompt[:200]}...")
    
    client = _get_client()
    if client is None:
        app.logger.error("GPT CLIENT IS NONE - OpenAI disabled or key missing")
        app.logger.error(f"OPENAI_API_KEY set: {bool(OPENAI_API_KEY)}")
        app.logger.error(f"OAI_ENABLED: {OAI_ENABLED}")
        app.logger.error(f"_openai_available: {_openai_available}")
        return "Placeholder output (model disabled or key missing)."
    
    try:
        app.logger.info(f"Sending request to OpenAI (timeout: {OAI_BUDGET}s)...")
        result = EXECUTOR.submit(_gpt_call, client, system, prompt).result(timeout=OAI_BUDGET)
        app.logger.info(f"OpenAI response length: {len(result)} chars")
        app.logger.info(f"OpenAI response preview: {result[:200]}...")
        humanized = humanize_text(result)
        app.logger.info("=== GPT CALL SUCCESS ===")
        return humanized
    except FuturesTimeout:
        app.logger.error(f"OpenAI request timed out after {OAI_BUDGET} seconds")
        return "Placeholder output due to timeout."
    except Exception as e:
        app.logger.error(f"OpenAI error: {type(e).__name__}: {e}")
        app.logger.error("=== GPT CALL FAILED ===")
        return "Placeholder output due to model error."

def extract_jd_keywords(jd: str, max_keywords: int = 8) -> List[str]:
    """Extract key skills and technologies from job description"""
    if not jd:
        return []
    
    # Common technical and business keywords to look for
    keyword_patterns = [
        r'\b(?:Python|Java|SQL|JavaScript|React|Node\.js|AWS|Azure|Docker|Kubernetes)\b',
        r'\b(?:Agile|Scrum|DevOps|CI/CD|ETL|API|REST|GraphQL|MongoDB|PostgreSQL)\b',
        r'\b(?:machine learning|data analysis|business intelligence|project management)\b',
        r'\b(?:stakeholder management|cross-functional|leadership|collaboration)\b',
        r'\b(?:analytics|optimization|automation|integration|scalability)\b'
    ]
    
    found_keywords = set()
    jd_lower = jd.lower()
    
    for pattern in keyword_patterns:
        matches = re.findall(pattern, jd_lower, re.IGNORECASE)
        found_keywords.update(matches)
    
    # Also look for key nouns (simple approach)
    important_nouns = ['team', 'product', 'system', 'platform', 'solution', 'process', 'strategy', 'performance']
    for noun in important_nouns:
        if noun in jd_lower:
            found_keywords.add(noun)
    
    return list(found_keywords)[:max_keywords]

# ----------------------------
# Enhanced GPT Bullets Generation
# ----------------------------
def gpt_bullets_batch(experience: List[Dict], jd: str, style_rules: List[str], metrics_by_company: Dict[str, List[str]]) -> Dict[str, List[str]]:
    """Generate humanized, quantified bullets - ENSURES EXACT COUNT"""
    app.logger.info("=== BULLETS GENERATION START ===")
    app.logger.info(f"Experience entries: {len(experience)}")
    app.logger.info(f"JD length: {len(jd)} chars")
    app.logger.info(f"JD preview: {jd[:300]}...")
    
    entries = []
    bullet_counts = {}
    
    for e in experience:
        k = int(e.get("bullets", 0) or 0)
        if k <= 0: continue
        comp = sanitize(e.get("company",""))
        role = sanitize(e.get("role",""))
        mx = metrics_by_company.get(comp, [])
        bullet_counts[comp] = k
        entries.append(f'- company: "{comp}"; role: "{role}"; bullets: {k}; metrics: [{", ".join(mx)}]')
        app.logger.info(f"Will generate {k} bullets for {comp} ({role}) with {len(mx)} metrics")
    
    if not entries:
        app.logger.warning("No experience entries to process")
        return {sanitize(e.get("company","")): [] for e in experience}

    # Extract key skills and requirements from job description
    jd_keywords = extract_jd_keywords(jd)
    app.logger.info(f"Extracted JD keywords: {jd_keywords}")

    # Enhanced style rules for human-like output
    rules_txt = "\n".join(f"- {r}" for r in (style_rules or [
        "20 to 28 words each bullet",
        "Start with strong action verb in past tense (Led, Developed, Implemented, Managed, etc.)",
        "Include specific numbers and percentages naturally using provided metrics",
        "Never use em/en dashes (–, —), use commas or connecting words like 'by', 'through', 'resulting in'",
        "Write like a human, avoid AI buzzwords and formal corporate speak",
        "Focus on measurable business impact and outcomes",
        "Use simple connecting words: 'and', 'by', 'through', 'resulting in', 'leading to'",
        "Each bullet must contain at least one quantified metric from the provided list",
        "Tailor content to match job requirements and keywords"
    ]))
    
    # More detailed prompt with explicit job description integration
    prompt = f"""You are a professional resume writer helping tailor resume bullets to a specific job posting. 

TARGET JOB REQUIREMENTS:
{jd[:2000]}

KEY SKILLS/KEYWORDS TO EMPHASIZE: {', '.join(jd_keywords)}

RESUME ENTRIES TO GENERATE (generate EXACTLY the number of bullets specified):
{chr(10).join(entries)}

STYLE REQUIREMENTS:
{rules_txt}

CRITICAL INSTRUCTIONS:
1. You MUST provide EXACTLY the number of bullets specified for each company. No more, no less.
2. Tailor each bullet to highlight skills/experience relevant to the job posting above
3. Use the provided metrics naturally within each bullet - do not invent new numbers
4. Write in a natural, human tone - avoid corporate buzzwords and AI-sounding language
5. Each bullet should demonstrate value delivered and impact achieved
6. Make bullets unique and specific to each role, not generic templates

Return ONLY valid JSON in this exact format:
{{
  "Company Name": ["bullet 1 with metrics", "bullet 2 with metrics", "...exactly as many as specified"]
}}

Example of good bullet style:
"Led cross-functional team of 8 developers to implement new CRM system, reducing customer response time by 40% and increasing satisfaction scores by 25%"

Generate the JSON now:"""

    app.logger.info(f"Final prompt length: {len(prompt)} chars")
    app.logger.info(f"JD keywords being emphasized: {jd_keywords}")
    
    text = gpt(prompt, system="You are a professional resume writer. Generate natural, human-sounding bullets tailored to the job requirements. Return only valid JSON with EXACTLY the requested number of bullets.")
    
    app.logger.info(f"Raw GPT response: {text[:500]}...")
    
    # More robust JSON parsing
    try:
        # Try to extract JSON from the response
        json_match = re.search(r'\{.*\}', text, re.DOTALL)
        if json_match:
            json_str = json_match.group(0)
            data = json.loads(json_str)
            app.logger.info(f"Successfully parsed JSON with {len(data)} companies")
        else:
            app.logger.error("No JSON found in response")
            data = {}
    except Exception as e:
        app.logger.error(f"JSON parsing failed: {e}")
        app.logger.error(f"Raw text was: {text}")
        data = {}
    
    # Ensure exact bullet count for each company
    out = {}
    for e in experience:
        comp = sanitize(e.get("company",""))
        k = int(e.get("bullets", 0) or 0)
        
        if k <= 0:
            out[comp] = []
            continue
        
        # Get bullets from GPT response
        gpt_bullets = [sanitize(humanize_text(b)) for b in (data.get(comp) or [])]
        
        # Ensure exactly k bullets
        if len(gpt_bullets) >= k:
            # If we have enough or more, take exactly k
            bullets = gpt_bullets[:k]
        else:
            # If we have fewer, generate fallback bullets to fill the gap
            bullets = gpt_bullets
            # Enhanced fallback bullets that use JD keywords
            jd_tailored_templates = [
                f"Enhanced {jd_keywords[0] if jd_keywords else 'operational'} efficiency by implementing data-driven solutions",
                f"Collaborated with stakeholders to deliver measurable improvements in {jd_keywords[1] if len(jd_keywords) > 1 else 'performance'} metrics",
                f"Drove {jd_keywords[2] if len(jd_keywords) > 2 else 'process'} improvements resulting in quantifiable business value",
                f"Managed complex initiatives focusing on {jd_keywords[0] if jd_keywords else 'strategic'} objectives and enhanced team productivity",
                f"Led strategic projects that delivered substantial ROI and improved business outcomes",
                f"Developed innovative approaches to challenges, achieving significant cost savings and efficiency gains"
            ]
            
            while len(bullets) < k:
                # Add a fallback bullet (cycle through templates)
                template_idx = (len(bullets) - len(gpt_bullets)) % len(jd_tailored_templates)
                fallback = jd_tailored_templates[template_idx]
                # Add some metrics if available
                if metrics_by_company.get(comp):
                    metric = metrics_by_company[comp][len(bullets) % len(metrics_by_company[comp])]
                    fallback = f"{fallback}, achieving {metric} improvement"
                bullets.append(sanitize(fallback))
        
        # Ensure quantification
        final_bullets = []
        for b in bullets:
            if not re.search(r'\d+', b):
                # Add a generic metric if none exists
                b = f"{b} achieving measurable results"
            final_bullets.append(b)
        
        out[comp] = final_bullets[:k]  # Ensure exactly k bullets
        app.logger.info(f"Generated {len(out[comp])} bullets for {comp}")
        
        # Log if we had to adjust
        if len(gpt_bullets) != k:
            app.logger.warning(f"GPT returned {len(gpt_bullets)} bullets for {comp}, expected {k}. Adjusted to match.")
    
    app.logger.info("=== BULLETS GENERATION COMPLETE ===")
    return out

# ----------------------------
# Metrics Extraction
# ----------------------------
def extract_numeric_phrases(text: str, max_phrases: int = 15) -> List[str]:
    """Extract quantified metrics from text"""
    NUM_REGEX = re.compile(
        r"(\$?\d+(?:\.\d+)?[KMB]?\b|\d+\+?\s*%|\d+\s*(?:hours?|days?|weeks?|months?|years?)|"
        r"\$\d{1,3}(?:,\d{3})*(?:\.\d+)?|\d{1,3}%|\d+x\b|\d+\s*million|\d+\s*billion)",
        re.IGNORECASE
    )
    
    found = []
    seen = set()
    for m in NUM_REGEX.finditer(text or ""):
        val = sanitize(m.group(0))
        if val and val.lower() not in seen:
            found.append(val)
            seen.add(val.lower())
            if len(found) >= max_phrases:
                break
    return found

# ----------------------------
# Debug and Test Routes
# ----------------------------
@app.route("/debug")
def debug():
    return jsonify({
        "openai_key_set": bool(OPENAI_API_KEY),
        "openai_key_length": len(OPENAI_API_KEY) if OPENAI_API_KEY else 0,
        "openai_available": _openai_available,
        "oai_enabled": OAI_ENABLED,
        "model": MODEL,
        "budget": OAI_BUDGET,
        "timeout": OAI_CLIENT_TIMEOUT,
        "threads": os.getenv("OAI_THREADS", "2"),
        "env_vars": {
            "OPENAI_MODEL": os.getenv("OPENAI_MODEL"),
            "USE_OPENAI": os.getenv("USE_OPENAI"),
            "OAI_BUDGET": os.getenv("OAI_BUDGET"),
            "OAI_CLIENT_TIMEOUT": os.getenv("OAI_CLIENT_TIMEOUT")
        }
    })

@app.route("/test-openai")
def test_openai():
    try:
        # Test basic GPT call
        test_result = gpt("Write exactly one sentence about software engineering.", 
                         "You are a helpful assistant.")
        
        # Test bullet generation with minimal data
        test_exp = [{"company": "Test Corp", "role": "Test Role", "bullets": 2}]
        test_jd = "We need a software engineer with Python experience."
        test_metrics = {"Test Corp": ["50%", "$1M", "100 users"]}
        
        test_bullets = gpt_bullets_batch(test_exp, test_jd, [], test_metrics)
        
        return jsonify({
            "success": True,
            "test_summary": test_result,
            "test_bullets": test_bullets,
            "summary_is_placeholder": "placeholder" in test_result.lower(),
            "bullets_is_placeholder": any("placeholder" in str(v).lower() for v in test_bullets.values())
        })
    
    except Exception as e:
        return jsonify({
            "success": False,
            "error": str(e),
            "error_type": type(e).__name__
        })

# ----------------------------
# Main Routes
# ----------------------------
@app.route("/")
def index():
    return jsonify({"ok": True, "service": "resume-tailor-server", "endpoints": ["/health", "/tailor", "/debug", "/test-openai"]})

@app.route("/health")
def health():
    enabled = bool(OPENAI_API_KEY) and OAI_ENABLED
    return jsonify({"ok": True, "model": MODEL, "openai_enabled": enabled})

@app.route("/tailor", methods=["POST","OPTIONS"])
def tailor():
    origin = request.headers.get("Origin", "*")
    
    # Parse request
    if request.content_type and "multipart/form-data" in request.content_type.lower():
        base_resume_file = request.files.get("base_resume")
        payload_part = request.form.get("payload")
        if not payload_part:
            return make_response(("missing payload", 400))
        try:
            data = json.loads(payload_part)
        except Exception:
            return make_response(("invalid payload json", 400))
    else:
        try:
            data = request.get_json(force=True, silent=False)
        except Exception:
            return make_response(("invalid json", 400))
        base_resume_file = None
    
    # Options
    global SHOW_KPI_PLACEHOLDER, BULLETS_STRICT_REPLACE
    BULLETS_STRICT_REPLACE = True  # Always replace bullets
    opts = data.get("options", {})
    
    exact_headers: Dict[str,str] = {}
    if "exact_headers" in opts and isinstance(opts["exact_headers"], dict):
        exact_headers = {sanitize(k): v for k, v in opts["exact_headers"].items() if v}
    
    # Inputs
    job_desc = sanitize(data.get("job_description",""))
    cfg = data.get("resume_config", {})
    summary_sentences = int(cfg.get("summary_sentences", 2))
    experience = cfg.get("experience", [])
    skills_categories = cfg.get("skills_categories", [])
    style_rules = opts.get("style_rules", [])
    
    # Load base resume
    if base_resume_file:
        base_doc = ensure_docx(base_resume_file)
    else:
        base_path = os.path.join(os.path.dirname(__file__), "base_resume.docx")
        if not os.path.exists(base_path):
            return make_response(("server missing base_resume.docx", 500))
        with open(base_path, "rb") as f:
            base_doc = ensure_docx(f)
    
    # Extract metrics from existing content BEFORE making any changes
    anchors_pre = find_anchors_by_exact_headers(base_doc, exact_headers) if exact_headers else []
    
    metrics_by_company: Dict[str, List[str]] = {}
    for comp, start_i in anchors_pre:
        # For metrics extraction, look at ALL content until next major section
        buf = []
        for i in range(start_i + 1, len(base_doc.paragraphs)):
            t = sanitize(base_doc.paragraphs[i].text)
            if is_major_heading(t):
                break
            if t:
                buf.append(t)
        
        role_metrics = extract_numeric_phrases(" ".join(buf))
        jd_metrics = extract_numeric_phrases(job_desc)
        
        merged = []
        seen = set()
        for val in role_metrics + jd_metrics:
            if val not in seen:
                seen.add(val)
                merged.append(val)
        
        metrics_by_company[comp] = merged[:20]
        app.logger.info(f"Extracted {len(role_metrics)} metrics for {comp}")
    
    # Generate professional summary
    summary_prompt = (
        f"Write exactly {summary_sentences} sentences for a professional summary. "
        f"Use natural language with specific achievements and numbers. "
        f"Avoid AI-style writing markers like dashes. Use 'and', 'through', 'by' as connectors.\n\n"
        f"Job Description:\n{job_desc[:1500]}"
    )
    summary = sanitize(gpt(summary_prompt))
    
    # Generate bullets
    bullets_by_company = gpt_bullets_batch(experience, job_desc, style_rules, metrics_by_company)
    
    # MODIFICATION 1 & 3: Replace summary completely
    replace_summary(base_doc, summary)
    
    # MODIFICATION 2: Preserve original skills, append new ones
    SKILL_BANK = {
        "Program & Project Delivery": [
            "Agile (Scrum, Kanban)","SDLC Management","Stakeholder Management","Risk Mitigation",
            "Change Management","Budget & Resource Planning","Vendor Management","Sprint planning"
        ],
        "Business & Systems Analysis": [
            "Requirements Engineering (BRD/FRD)","User Stories","Business Process Modeling (BPMN)",
            "Go-to-Market (GTM) Strategy","Data Governance (GDPR, CCPA)","Process re-engineering","UAT"
        ],
        "AI & Data Analytics": [
            "Generative AI (OpenAI API, RAG)","Predictive Modeling (Python, scikit-learn, XGBoost)",
            "Data Visualization (Power BI, Tableau)","SQL","Natural Language Processing (NLP)","Machine Learning"
        ],
        "Cloud & Enterprise Platforms": [
            "AWS (Lambda, S3, Redshift, EMR)","SAP S/4HANA","Frappe ERP","Jira","Confluence",
            "REST APIs","CI/CD Pipelines","Docker","Kubernetes"
        ],
        "Industry Expertise": [
            "HealthTech","Enterprise SaaS","Consulting","HR Technology","Supply Chain",
            "Logistics","E-commerce & Retail","Financial Services"
        ]
    }
    
    def pick_relevant_skills(jd: str, bank: Dict[str, List[str]], top_k: int = 6) -> Dict[str, List[str]]:
        jd_lower = jd.lower()
        out = {}
        for cat, items in bank.items():
            relevant = []
            for skill in items:
                pattern = re.escape(skill.lower()).replace(r"\ ", r"\s*")
                if re.search(rf"\b{pattern}\b", jd_lower):
                    relevant.append(skill)
            
            # Take relevant skills first, then fill with others
            seen = set()
            merged = []
            for s in relevant + items:
                if s not in seen:
                    seen.add(s)
                    merged.append(s)
                    if len(merged) >= top_k:
                        break
            
            if merged:
                out[cat] = merged
        
        return out
    
    skills_categories = skills_categories or list(SKILL_BANK.keys())
    skills_map = pick_relevant_skills(job_desc, {c: SKILL_BANK.get(c, []) for c in skills_categories})
    merge_skills(base_doc, skills_map)
    
    # MODIFICATION 1: Replace bullets - ROBUST VERSION
    # Build exact_headers if not provided
    if not exact_headers and experience:
        exact_headers = {}
        for exp in experience:
            comp = sanitize(exp.get("company", ""))
            role = sanitize(exp.get("role", ""))
            if comp and role:
                # Try to find the header in the document that contains both company and role
                for para in base_doc.paragraphs:
                    para_text = sanitize(para.text)
                    if comp in para_text and role in para_text:
                        exact_headers[comp] = para_text
                        app.logger.info(f"Found header for {comp}: {para_text[:80]}")
                        break
    
    # Re-find anchors after summary/skills changes
    anchors = find_anchors_by_exact_headers(base_doc, exact_headers) if exact_headers else []
    
    if anchors:
        app.logger.info(f"\n{'='*60}")
        app.logger.info(f"Starting bullet replacement for {len(anchors)} roles")
        
        # Process each company one at a time
        for comp, _ in anchors:
            app.logger.info(f"\nProcessing {comp}")
            
            # Find the header for this company (fresh search each time)
            header_idx = None
            header_text = exact_headers.get(comp, "")
            
            if header_text:
                for i, para in enumerate(base_doc.paragraphs):
                    if sanitize(para.text) == sanitize(header_text):
                        header_idx = i
                        break
            
            # Fallback: search for company name in paragraph
            if header_idx is None:
                for i, para in enumerate(base_doc.paragraphs):
                    if comp in sanitize(para.text):
                        # Check if this looks like a role header (has bold or is in professional experience section)
                        has_bold = any(run.bold for run in para.runs) if para.runs else False
                        if has_bold or "analyst" in para.text.lower() or "manager" in para.text.lower():
                            header_idx = i
                            header_text = para.text
                            app.logger.info(f"Found header via fallback search: {header_text[:80]}")
                            break
            
            if header_idx is None:
                app.logger.warning(f"Could not find header for {comp}")
                continue
            
            app.logger.info(f"Found {comp} header at index {header_idx}: {base_doc.paragraphs[header_idx].text[:80]}")
            
            # Find all content to delete (everything until next header or section)
            indices_to_delete = []
            i = header_idx + 1
            
            while i < len(base_doc.paragraphs):
                para_text = base_doc.paragraphs[i].text.strip()
                
                # Stop if we hit a major section
                if is_major_heading(para_text):
                    app.logger.info(f"  Stopping at major heading: {para_text[:50]}")
                    break
                
                # Check if this is another role header
                para = base_doc.paragraphs[i]
                has_bold = any(run.bold for run in para.runs) if para.runs else False
                
                # If it has bold and contains a company name (not current company)
                if has_bold and para_text:
                    # Check if this is a different company's header
                    is_other_company = False
                    for other_comp in exact_headers.keys():
                        if other_comp != comp and other_comp in para_text:
                            is_other_company = True
                            break
                    
                    # Also check for role keywords to confirm it's a header
                    role_keywords = ["analyst", "manager", "engineer", "developer", "consultant", "director"]
                    has_role = any(keyword in para_text.lower() for keyword in role_keywords)
                    
                    if is_other_company and has_role:
                        app.logger.info(f"  Stopping at next company header: {para_text[:50]}")
                        break
                
                # This paragraph belongs to current role - mark for deletion
                indices_to_delete.append(i)
                if para_text:
                    app.logger.info(f"  Will delete at {i}: {para_text[:80]}")
                
                i += 1
            
            # Delete all marked paragraphs in reverse order
            if indices_to_delete:
                app.logger.info(f"Deleting {len(indices_to_delete)} paragraphs for {comp}")
                for idx in reversed(indices_to_delete):
                    try:
                        p = base_doc.paragraphs[idx]._element
                        p.getparent().remove(p)
                    except Exception as e:
                        app.logger.warning(f"Could not delete paragraph at {idx}: {e}")
            
            # Find the header again (indices have changed)
            header_idx_new = None
            for i, para in enumerate(base_doc.paragraphs):
                if comp in sanitize(para.text):
                    # Verify this is the right header
                    has_bold = any(run.bold for run in para.runs) if para.runs else False
                    if has_bold or "analyst" in para.text.lower() or "manager" in para.text.lower():
                        header_idx_new = i
                        break
            
            if header_idx_new is None:
                app.logger.error(f"Could not re-find header for {comp}")
                continue
            
            # Insert new bullets
            last = base_doc.paragraphs[header_idx_new]
            new_bullets = bullets_by_company.get(comp, [])
            
            app.logger.info(f"Inserting {len(new_bullets)} new bullets for {comp} after index {header_idx_new}")
            for bullet in new_bullets:
                bullet_text = sanitize(bullet)
                if not bullet_text.startswith("•"):
                    bullet_text = f"• {bullet_text}"
                new_para = insert_paragraph_after(last, bullet_text)
                set_paragraph_font(new_para)
                last = new_para
        
        app.logger.info(f"\n{'='*60}")
        app.logger.info(f"Bullet replacement complete")
    else:
        app.logger.warning("No anchors found - attempting direct search")
        # Fallback: try to process based on experience config alone
        for exp in experience:
            comp = sanitize(exp.get("company", ""))
            if not comp:
                continue
                
            app.logger.info(f"\nProcessing {comp} via direct search")
            
            # Find header
            header_idx = None
            for i, para in enumerate(base_doc.paragraphs):
                if comp in sanitize(para.text):
                    has_bold = any(run.bold for run in para.runs) if para.runs else False
                    if has_bold:
                        header_idx = i
                        app.logger.info(f"Found {comp} at index {i}")
                        break
            
            if header_idx is None:
                continue
            
            # Delete old content and insert new (similar logic as above)
            # ... (would continue with deletion and insertion logic)
    
    # MODIFICATION 4: Generate filename WITHOUT "Tailored"
    today = datetime.now().strftime("%Y-%m-%d")
    filename = f"Rijul_Chaturvedi_{today}.docx"
    
    # Save and send
    out = io.BytesIO()
    base_doc.save(out)
    out.seek(0)
    
    resp = make_response(send_file(
        out,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=filename,
    ))
    resp.headers["Access-Control-Allow-Origin"] = origin
    resp.headers["Vary"] = "Origin"
    return resp

if __name__ == "__main__":
    port = int(os.getenv("PORT", "8000"))
    app.run(host="0.0.0.0", port=port)